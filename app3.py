# import streamlit as st
# import openai
# import docx
# import os
#
# # ─── PAGE CONFIG ───────────────────────────────────────────────────────────
# st.set_page_config(page_title="📚 Ask from DOCX Files", layout="wide")
#
# # ─── SIDEBAR ───────────────────────────────────────────────────────────────
# st.sidebar.header("🔐 API & Model Settings")
# api_key = st.sidebar.text_input("OpenAI API Key", type="password")
# model = st.sidebar.selectbox("Model", ["gpt-4o", "gpt-3.5-turbo"])
#
# st.sidebar.markdown("---")
# st.sidebar.header("✂️ Response Limit Setting")
#
# limit_type = st.sidebar.radio("Select limit type", ["Word Limit", "Character Limit"])
# if limit_type == "Word Limit":
#     word_limit = st.sidebar.number_input("Set Word Limit", min_value=10, max_value=1000, value=150)
#     char_limit = None
# else:
#     char_limit = st.sidebar.number_input("Set Character Limit", min_value=50, max_value=4000, value=1000)
#     word_limit = None
#
# # ─── FILE UPLOAD ───────────────────────────────────────────────────────────
# st.header("📁 Upload DOCX Files and Ask a Question")
# uploaded_files = st.file_uploader("Upload two or more Word documents", type=["docx"], accept_multiple_files=True)
#
# question = st.text_area("💬 Enter your question", placeholder="e.g., What are the key points discussed in the sessions?")
#
# submit = st.button("🤖 Generate Answer")
#
# # ─── FUNCTION: Extract Text from DOCX ──────────────────────────────────────
# def extract_text_from_docx(uploaded_files):
#     all_text = []
#     for file in uploaded_files:
#         try:
#             doc = docx.Document(file)
#             content = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
#             all_text.append(content)
#         except Exception as e:
#             st.error(f"Error reading {file.name}: {e}")
#     return "\n\n".join(all_text)
#
# # ─── FUNCTION: Ask OpenAI ──────────────────────────────────────────────────
# def ask_openai(question, context, model, api_key, word_limit=None, char_limit=None):
#     openai.api_key = api_key
#
#     limit_description = ""
#     if word_limit:
#         limit_description = f"Keep the answer under {word_limit} words."
#     elif char_limit:
#         limit_description = f"Keep the answer under {char_limit} characters."
#
#     system_prompt = f"""
# You are a helpful assistant. Answer the user's question using the content provided below.
#
# Only use the uploaded document content as your source of truth. If something is unclear, do your best to interpret it based on the available information. Avoid adding outside knowledge.
#
# {limit_description}
#
# Content:
# {context}
# """
#
#     try:
#         response = openai.chat.completions.create(
#             model=model,
#             messages=[
#                 {"role": "system", "content": system_prompt.strip()},
#                 {"role": "user", "content": question.strip()}
#             ],
#             max_tokens=4096,
#             temperature=0.7
#         )
#         return response.choices[0].message.content.strip()
#     except Exception as e:
#         return f"❌ Error calling OpenAI API: {e}"
#
# # ─── MAIN EXECUTION ────────────────────────────────────────────────────────
# if submit:
#     if not api_key:
#         st.error("🔐 Please enter your OpenAI API key.")
#     elif not uploaded_files or len(uploaded_files) < 2:
#         st.warning("📂 Please upload at least two .docx files.")
#     elif not question.strip():
#         st.warning("❓ Please enter a question to ask.")
#     else:
#         with st.spinner("Extracting text from documents..."):
#             combined_text = extract_text_from_docx(uploaded_files)
#
#         with st.spinner("Sending to OpenAI..."):
#             answer = ask_openai(question, combined_text, model, api_key, word_limit, char_limit)
#
#         st.success("✅ Answer Generated")
#         st.text_area("📘 Answer", answer, height=300)

import streamlit as st
import openai
import docx
import pandas as pd
import io
import time

# ─── PAGE CONFIG ───────────────────────────────────────────────────────────
st.set_page_config(page_title="📚 Bulk QA from DOCX + Excel", layout="wide")

# ─── SIDEBAR SETTINGS ──────────────────────────────────────────────────────
st.sidebar.header("🔐 API & Model Settings")
api_key = st.sidebar.text_input("OpenAI API Key", type="password")
model = st.sidebar.selectbox("Model", ["gpt-4o", "gpt-3.5-turbo"])

st.sidebar.markdown("---")
st.sidebar.header("🧠 System Prompt Template")

default_prompt = """
You are a helpful assistant. Answer the user's question using the content provided below.

Only use the uploaded document content as your source of truth. If something is unclear, do your best to interpret it based on the available information. Avoid adding outside knowledge.

Your answer must be exactly {limit_value} {limit_type} long.

Content:
{context}
"""

user_prompt_template = st.sidebar.text_area(
    "Edit System Prompt (use {limit_value}, {limit_type}, and {context})",
    value=default_prompt.strip(),
    height=300
)

# ─── FILE UPLOAD ───────────────────────────────────────────────────────────
st.header("📁 Upload DOCX Files and Excel with Questions")
uploaded_files = st.file_uploader("Upload DOCX Files", type=["docx"], accept_multiple_files=True)
excel_file = st.file_uploader("Upload Excel File (.xlsx)", type=["xlsx"])

submit = st.button("🤖 Generate Answers")

# ─── FUNCTION: Extract Text from DOCX ──────────────────────────────────────
def extract_text_from_docx(files):
    text_parts = []
    for file in files:
        try:
            doc = docx.Document(file)
            content = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
            text_parts.append(content)
        except Exception as e:
            st.error(f"Error reading {file.name}: {e}")
    return "\n\n".join(text_parts)

# ─── FUNCTION: Ask OpenAI ──────────────────────────────────────────────────
def ask_openai(question, context, model, api_key, limit_value, limit_type, system_prompt_template):
    openai.api_key = api_key

    prompt_filled = system_prompt_template.replace("{limit_value}", str(limit_value)) \
                                          .replace("{limit_type}", limit_type) \
                                          .replace("{context}", context)

    try:
        response = openai.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": prompt_filled.strip()},
                {"role": "user", "content": question.strip()}
            ],
            max_tokens=4096,
            temperature=0.7
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"❌ Error: {e}"

# ─── MAIN EXECUTION ────────────────────────────────────────────────────────
if submit:
    if not api_key:
        st.error("🔐 Please enter your OpenAI API key.")
    elif not uploaded_files or len(uploaded_files) < 1:
        st.warning("📂 Please upload at least one DOCX file.")
    elif not excel_file:
        st.warning("📄 Please upload an Excel file.")
    else:
        with st.spinner("📖 Extracting text from documents..."):
            doc_text = extract_text_from_docx(uploaded_files)

        try:
            df = pd.read_excel(excel_file)
        except Exception as e:
            st.error(f"Failed to read Excel file: {e}")
            st.stop()

        if "Question" not in df.columns or ("Word Limit" not in df.columns and "Character Limit" not in df.columns):
            st.error("Excel must contain columns: 'Question', and either 'Word Limit' or 'Character Limit'.")
            st.stop()

        st.info(f"📋 Found {len(df)} questions. Generating answers...")

        answers = []
        progress_bar = st.progress(0)
        status_text = st.empty()

        for i, row in df.iterrows():
            question = str(row["Question"])
            word_limit = row.get("Word Limit")
            char_limit = row.get("Character Limit")

            if pd.notna(word_limit):
                limit_value = int(word_limit)
                limit_type = "words"
            elif pd.notna(char_limit):
                limit_value = int(char_limit)
                limit_type = "characters"
            else:
                answers.append("❌ No limit specified.")
                continue

            status_text.text(f"Processing question {i+1}/{len(df)}...")
            answer = ask_openai(question, doc_text, model, api_key, limit_value, limit_type, user_prompt_template)
            answers.append(answer)

            progress_bar.progress((i + 1) / len(df))
            time.sleep(0.5)  # Slight delay for visual progress

        df["Answer"] = answers

        st.success("✅ All answers generated!")
        st.dataframe(df)

        # Download output
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name="Answers")
        st.download_button("📥 Download Results as Excel", data=output.getvalue(), file_name="answers.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

